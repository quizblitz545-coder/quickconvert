import os
import tempfile
from PIL import Image
import PyPDF2
from docx import Document
import markdown
import weasyprint
from pydub import AudioSegment

class DocumentConverter:
    """Handles document format conversions"""
    
    def convert(self, source_path, output_path, target_format):
        """Convert document to target format"""
        try:
            source_ext = os.path.splitext(source_path)[1].lower()
            
            # PDF conversions
            if source_ext == '.pdf':
                return self._convert_from_pdf(source_path, output_path, target_format)
            
            # DOCX conversions
            elif source_ext == '.docx':
                return self._convert_from_docx(source_path, output_path, target_format)
            
            # TXT conversions
            elif source_ext == '.txt':
                return self._convert_from_txt(source_path, output_path, target_format)
            
            # HTML conversions
            elif source_ext == '.html':
                return self._convert_from_html(source_path, output_path, target_format)
            
            # Markdown conversions
            elif source_ext == '.md':
                return self._convert_from_markdown(source_path, output_path, target_format)
            
            return False
            
        except Exception as e:
            print(f"Document conversion error: {e}")
            return False
    
    def _convert_from_pdf(self, source_path, output_path, target_format):
        """Convert PDF to other formats"""
        try:
            with open(source_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text_content = ""
                
                for page in pdf_reader.pages:
                    text_content += page.extract_text() + "\n"
            
            if target_format == 'txt':
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(text_content)
                return True
            
            elif target_format == 'html':
                html_content = f"<html><body><pre>{text_content}</pre></body></html>"
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(html_content)
                return True
            
            elif target_format == 'docx':
                doc = Document()
                doc.add_paragraph(text_content)
                doc.save(output_path)
                return True
                
        except Exception as e:
            print(f"PDF conversion error: {e}")
            return False
    
    def _convert_from_docx(self, source_path, output_path, target_format):
        """Convert DOCX to other formats"""
        try:
            doc = Document(source_path)
            text_content = ""
            
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            if target_format == 'txt':
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(text_content)
                return True
            
            elif target_format == 'html':
                html_content = f"<html><body><pre>{text_content}</pre></body></html>"
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(html_content)
                return True
            
            elif target_format == 'pdf':
                html_content = f"<html><body><pre>{text_content}</pre></body></html>"
                weasyprint.HTML(string=html_content).write_pdf(output_path)
                return True
                
        except Exception as e:
            print(f"DOCX conversion error: {e}")
            return False
    
    def _convert_from_txt(self, source_path, output_path, target_format):
        """Convert TXT to other formats"""
        try:
            with open(source_path, 'r', encoding='utf-8') as f:
                text_content = f.read()
            
            if target_format == 'html':
                html_content = f"<html><body><pre>{text_content}</pre></body></html>"
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(html_content)
                return True
            
            elif target_format == 'pdf':
                html_content = f"<html><body><pre>{text_content}</pre></body></html>"
                weasyprint.HTML(string=html_content).write_pdf(output_path)
                return True
            
            elif target_format == 'docx':
                doc = Document()
                doc.add_paragraph(text_content)
                doc.save(output_path)
                return True
                
        except Exception as e:
            print(f"TXT conversion error: {e}")
            return False
    
    def _convert_from_html(self, source_path, output_path, target_format):
        """Convert HTML to other formats"""
        try:
            with open(source_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            
            if target_format == 'pdf':
                weasyprint.HTML(filename=source_path).write_pdf(output_path)
                return True
            
            elif target_format == 'txt':
                # Simple HTML to text conversion (basic)
                import re
                text_content = re.sub('<[^<]+?>', '', html_content)
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(text_content)
                return True
                
        except Exception as e:
            print(f"HTML conversion error: {e}")
            return False
    
    def _convert_from_markdown(self, source_path, output_path, target_format):
        """Convert Markdown to other formats"""
        try:
            with open(source_path, 'r', encoding='utf-8') as f:
                md_content = f.read()
            
            if target_format == 'html':
                html_content = markdown.markdown(md_content)
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(html_content)
                return True
            
            elif target_format == 'pdf':
                html_content = markdown.markdown(md_content)
                weasyprint.HTML(string=html_content).write_pdf(output_path)
                return True
                
        except Exception as e:
            print(f"Markdown conversion error: {e}")
            return False


class ImageConverter:
    """Handles image format conversions"""
    
    def convert(self, source_path, output_path, target_format):
        """Convert image to target format"""
        try:
            with Image.open(source_path) as img:
                # Handle transparency for formats that don't support it
                if target_format.lower() in ['jpg', 'jpeg'] and img.mode in ('RGBA', 'LA', 'P'):
                    # Create white background
                    background = Image.new('RGB', img.size, (255, 255, 255))
                    if img.mode == 'P':
                        img = img.convert('RGBA')
                    background.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
                    img = background
                
                # Convert and save
                if target_format.lower() == 'jpg':
                    img.save(output_path, 'JPEG', quality=95)
                else:
                    img.save(output_path, target_format.upper())
                
                return True
                
        except Exception as e:
            print(f"Image conversion error: {e}")
            return False


class AudioConverter:
    """Handles audio format conversions"""
    
    def convert(self, source_path, output_path, target_format):
        """Convert audio to target format"""
        try:
            # Load audio file
            audio = AudioSegment.from_file(source_path)
            
            # Export to target format
            if target_format == 'mp3':
                audio.export(output_path, format="mp3")
            elif target_format == 'wav':
                audio.export(output_path, format="wav")
            elif target_format == 'ogg':
                audio.export(output_path, format="ogg")
            elif target_format == 'm4a':
                audio.export(output_path, format="mp4")
            else:
                return False
            
            return True
            
        except Exception as e:
            print(f"Audio conversion error: {e}")
            return False

